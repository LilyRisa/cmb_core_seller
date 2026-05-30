# -*- coding: utf-8 -*-
"""
Sinh tài liệu thiết kế / PRD cho sản phẩm CMBcoreSeller ra file Word (.docx).

Phiên bản dành cho NGƯỜI KIỂM DUYỆT: mô tả luồng dữ liệu & hệ thống bằng ngôn ngữ
nghiệp vụ thuần — KHÔNG chứa đường dẫn tài liệu nội bộ, tên hàm/biến/bảng/sự kiện,
mã route hay mã trạng thái kỹ thuật. Chạy:  python scripts/generate_prd.py
Yêu cầu: pip install python-docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs", "PRD-CMBcoreSeller.docx")
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_update_fields(doc):
    el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true")
    doc.settings.element.append(el)


def add_toc(doc):
    p = doc.add_paragraph(); run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Mục lục — mở trong Word rồi nhấn Ctrl+A, F9 để cập nhật."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, instr, f2, t, f3):
        run._r.append(el)


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY
    return p


def para(doc, text, italic=False, bold=False, size=None, color=None, space_after=6):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.italic, r.bold = italic, bold
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def flow(doc, text):
    """Một dòng tóm tắt luồng bằng mũi tên, ngôn ngữ thuần."""
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Luồng:  " + text); r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(10)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, (list, tuple)):
            r = p.add_run(it[0] + ": "); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""; run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.size = Pt(9.5)
        shade(hdr[i], "1F3A5F")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2); sec.top_margin = sec.bottom_margin = Cm(2)
cp = doc.core_properties
cp.title = "CMBcoreSeller — Tài liệu thiết kế (PRD)"; cp.author = "CMBcoreSeller"
cp.subject = "Product Requirements / Design Document"; cp.category = "PRD"

# Title page
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER; tp.paragraph_format.space_before = Pt(80)
r = tp.add_run("CMBcoreSeller"); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = NAVY
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Tài liệu Thiết kế Sản phẩm (PRD)"); r.font.size = Pt(16); r.font.color.rgb = GREY
sb = doc.add_paragraph(); sb.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sb.add_run("Nền tảng SaaS quản lý bán hàng đa sàn cho thị trường Việt Nam"); r.italic = True; r.font.size = Pt(12)
sc = doc.add_paragraph(); sc.alignment = WD_ALIGN_PARAGRAPH.CENTER; sc.paragraph_format.space_before = Pt(24)
r = sc.add_run("Mô tả cách hệ thống hoạt động: luồng dữ liệu · tính năng chính · trường hợp sử dụng"); r.font.size = Pt(11); r.font.color.rgb = NAVY
doc.add_page_break()

# Metadata + TOC
h(doc, "Thông tin tài liệu", 1)
table(doc, ["Mục", "Nội dung"], [
    ["Sản phẩm", "CMBcoreSeller — SaaS quản lý bán hàng đa sàn cho thị trường Việt Nam"],
    ["Loại tài liệu", "Tài liệu thiết kế / Mô tả yêu cầu sản phẩm (PRD)"],
    ["Phiên bản", "1.0"],
    ["Ngày", "2026-05-26"],
    ["Phạm vi", "Toàn bộ sản phẩm"],
    ["Mục đích", "Giúp người kiểm duyệt hiểu luồng dữ liệu và cách hệ thống vận hành"],
], widths=[4, 13])
h(doc, "Mục lục", 1)
add_toc(doc)
doc.add_page_break()

# 1. Tổng quan
h(doc, "1. Tổng quan sản phẩm", 1)
para(doc, "CMBcoreSeller là phần mềm dịch vụ trực tuyến (SaaS) giúp người bán hàng tại Việt Nam kết nối nhiều "
          "gian hàng trên TikTok Shop, Shopee, Lazada (và tạo đơn thủ công), gom đơn hàng về một nơi với trạng "
          "thái thống nhất, đồng bộ tồn kho theo từng mã hàng, xử lý giao hàng và in vận đơn hàng loạt, đồng thời "
          "quản lý kho – sản phẩm – tài chính – kế toán – chăm sóc khách hàng trên cùng một hệ thống.")
h(doc, "1.1 Đối tượng người dùng", 2)
bullets(doc, [
    ("Nhà bán hàng đa sàn", "một người quản lý nhiều gian hàng trên nhiều sàn."),
    ("Trong một nhà bán", "chủ shop và nhân viên (kho, xử lý đơn, kế toán, chăm sóc khách hàng) với phân quyền khác nhau."),
    ("Quy mô mục tiêu", "khoảng 100 nhà bán, mỗi nhà khoảng 5.000 đơn mỗi tháng."),
])
h(doc, "1.2 Vai trò & phân quyền", 2)
table(doc, ["Vai trò", "Phạm vi chính"], [
    ["Chủ shop / Quản trị", "Toàn quyền trong nhà bán: kết nối sàn, cấu hình, gói cước, quản lý nhân sự."],
    ["Nhân viên xử lý đơn", "Xử lý đơn, ghép mã hàng, tạo vận đơn, in tem."],
    ["Nhân viên kho", "Nhập, xuất, điều chuyển kho, kiểm kê, theo dõi giá vốn."],
    ["Kế toán", "Tài chính, đối soát, kế toán (ghi sổ, công nợ, báo cáo tài chính)."],
    ["Nhân viên chăm sóc khách hàng", "Hộp thư hợp nhất, trả lời tin nhắn, sử dụng trợ lý AI."],
    ["Người xem", "Chỉ xem, không chỉnh sửa."],
    ["Quản trị hệ thống", "Nhà vận hành nền tảng: quản lý gói cước, ưu đãi, nhà cung cấp AI (tách biệt với nhà bán)."],
], widths=[5, 12])

# 2. Tính năng chính
h(doc, "2. Các tính năng chính", 1)
para(doc, "Bộ tính năng đầy đủ cho hoạt động bán hàng đa sàn, chia theo nhóm nghiệp vụ.", italic=True)
table(doc, ["Nhóm tính năng", "Mô tả chính"], [
    ["Đơn hàng", "Đồng bộ đơn đa sàn (tức thời + định kỳ); tạo đơn thủ công; trạng thái thống nhất; lọc/tìm; thao tác hàng loạt; cảnh báo đơn trùng/lỗi/hết hàng; lợi nhuận ước tính."],
    ["Gian hàng", "Kết nối/ngắt kết nối gian hàng; nhật ký đồng bộ và cho chạy lại; tự gia hạn quyền truy cập và cảnh báo khi cần kết nối lại."],
    ["Tồn kho & quản lý kho", "Mã hàng gốc; nhiều kho; theo dõi tồn thực có / đang giữ / tồn an toàn / khả dụng; sổ biến động tồn; nhập–xuất–điều chuyển–kiểm kê có phiếu và duyệt; giá vốn theo FIFO và bình quân; chống bán vượt."],
    ["Sản phẩm & đăng bán", "Sản phẩm và mã hàng gốc; ghép mã hàng (một-một và combo nhiều mã); đồng bộ tồn hai chiều; đăng bán nhiều sàn từ một sản phẩm gốc."],
    ["Giao hàng & in ấn", "Tạo vận đơn (qua đơn vị vận chuyển của sàn hoặc của riêng nhà bán); tải tem thật; in vận đơn hàng loạt thành một tệp; tự tạo phiếu soạn hàng / đóng gói; quét mã đóng gói; lưu và in lại trong 90 ngày."],
    ["Sổ khách hàng", "Khớp đơn theo số điện thoại để dựng hồ sơ khách + thống kê + mức độ tín nhiệm; ghi chú tự động; chặn khách; ẩn danh hoá theo yêu cầu của sàn."],
    ["Mua hàng", "Nhà cung cấp và bảng giá nhập; đề xuất số lượng cần nhập; tạo đơn mua → nhận hàng → nhập kho → cập nhật giá vốn."],
    ["Tài chính", "Kéo bảng đối soát phí từ từng sàn; tính lợi nhuận theo đơn / sản phẩm / gian hàng / thời gian dựa trên phí thực tế."],
    ["Báo cáo", "Bảng điều khiển tổng quan; doanh thu / lợi nhuận / sản phẩm bán chạy; xuất tệp Excel/CSV."],
    ["Kế toán", "Ghi sổ kép theo chuẩn mực kế toán Việt Nam (TT133); hệ thống tài khoản; kỳ kế toán và khoá kỳ; ghi sổ tự động theo nghiệp vụ; công nợ phải thu/phải trả; quỹ & ngân hàng; thuế GTGT; báo cáo tài chính và xuất tệp tương thích phần mềm kế toán (MISA)."],
    ["Nhắn tin & trợ lý AI", "Hộp thư hợp nhất đa sàn và Facebook; mẫu trả lời; trả lời tự động theo điều kiện; trợ lý AI gợi ý hoặc tự trả lời qua bộ kiểm soát ý định; hỏi-đáp dựa trên tài liệu của shop; cho phép cấu hình nhà cung cấp AI tuỳ chỉnh."],
    ["Nền tảng & vận hành", "Nhiều nhà bán độc lập; tài khoản phụ và phân quyền; nhật ký thao tác; gói thuê bao + hạn mức + dùng thử + cổng thanh toán Việt Nam; công cụ vận hành cho quản trị hệ thống."],
])

# 3. Nguyên tắc thiết kế hệ thống
h(doc, "3. Nguyên tắc thiết kế hệ thống", 1)
para(doc, "Các nguyên tắc dưới đây giúp hiểu phần Luồng dữ liệu ở mục 4.")
bullets(doc, [
    ("Nhiều nhà bán độc lập", "mỗi nhà bán là một không gian làm việc riêng; dữ liệu được cách ly hoàn toàn giữa các nhà bán."),
    ("Tích hợp dạng mô-đun cắm-thêm", "mỗi sàn, đơn vị vận chuyển, cổng thanh toán, kênh nhắn tin và mô hình AI là một mô-đun độc lập; phần lõi không phụ thuộc vào tên nhà cung cấp cụ thể, nên thêm nhà cung cấp mới không ảnh hưởng phần còn lại."),
    ("Một nguồn sự thật", "tồn kho luôn tính theo mã hàng gốc; trạng thái đơn dùng một bộ trạng thái chuẩn chung, được ánh xạ từ trạng thái riêng của mỗi sàn."),
    ("Đồng bộ tin cậy", "nhận dữ liệu tức thời từ sàn, đồng thời đồng bộ định kỳ để không bỏ sót; mọi thao tác đồng bộ đều an toàn khi lặp lại nên không nhân đôi đơn hay tồn kho."),
    ("Xử lý nền", "các tác vụ nặng (đồng bộ, tạo tệp PDF, gửi tin) chạy ở chế độ nền và tự thử lại khi lỗi tạm thời."),
    ("Tiền tệ & thời gian", "chỉ dùng đồng Việt Nam (VND); mốc thời gian theo chuẩn quốc tế."),
])

# 4. LUỒNG DỮ LIỆU
h(doc, "4. Luồng dữ liệu chính", 1)
para(doc, "Phần trọng tâm. Mỗi luồng gồm: tóm tắt một dòng, các bước xử lý, và những điểm đảm bảo.")

h(doc, "4.1 Kết nối gian hàng", 2)
flow(doc, "Người bán bấm Kết nối → uỷ quyền trên sàn → hệ thống lưu kết nối → kéo lịch sử đơn + nhận thông báo từ sàn")
numbered(doc, [
    "Người bán chọn sàn cần kết nối và bấm “Kết nối”.",
    "Hệ thống chuyển người bán sang trang uỷ quyền của sàn để đăng nhập và đồng ý cấp quyền.",
    "Sàn xác nhận; hệ thống lưu kết nối của gian hàng (thông tin truy cập được mã hoá).",
    "Hệ thống tự kéo về lịch sử đơn khoảng 90 ngày và đăng ký nhận thông báo thay đổi từ sàn.",
    "Quyền truy cập được tự gia hạn định kỳ; nếu không gia hạn được, hệ thống báo người bán kết nối lại.",
])

h(doc, "4.2 Đồng bộ đơn hàng", 2)
flow(doc, "Sàn báo / hệ thống kéo định kỳ → lấy chi tiết đơn từ sàn → lưu đơn (chuẩn hoá trạng thái) → tự động giữ tồn · khớp khách · thông báo")
para(doc, "Hai chiều bổ trợ nhau: (a) sàn chủ động báo khi có thay đổi (tức thời); (b) hệ thống chủ động kéo định kỳ "
          "để phòng khi sót — gồm kéo theo thời gian, kéo toàn bộ khi mới kết nối, và kéo các đơn “chưa xử lý” "
          "(đặt lâu chưa giao).")
numbered(doc, [
    "Khi có đơn mới hoặc đổi trạng thái, hệ thống luôn lấy chi tiết đơn trực tiếp từ sàn (không tin dữ liệu thông báo thô).",
    "Lưu hoặc cập nhật đơn một cách an toàn-khi-lặp-lại: bỏ qua nếu dữ liệu cũ hơn bản đã có; ánh xạ trạng thái của sàn sang trạng thái chuẩn; ghi lại lịch sử trạng thái.",
    "Sau khi lưu đơn, hệ thống tự động: giữ tồn cho đơn, khớp đơn vào hồ sơ khách theo số điện thoại, chạy quy tắc tự động và gửi thông báo nếu được cấu hình.",
])
bullets(doc, [
    ("Đảm bảo", "không trùng đơn; một gian hàng lỗi không làm dừng gian hàng khác; tôn trọng giới hạn tần suất gọi của sàn; có nhật ký để xem lại và chạy lại khi cần."),
])

h(doc, "4.3 Tồn kho, ghép mã hàng & đẩy tồn lên sàn", 2)
flow(doc, "Ghép mã hàng sàn ↔ mã gốc → vòng đời đơn giữ/nhả/trừ tồn → tính lại tồn khả dụng → tự cập nhật lên sàn")
bullets(doc, [
    ("Ghép mã hàng", "khi đồng bộ sản phẩm từ sàn, hệ thống tự gợi ý ghép mã hàng trên sàn với mã hàng gốc khi trùng mã; mã chưa ghép được đánh dấu để người bán xử lý, đơn chứa mã chưa ghép được gắn cờ cần xử lý."),
    ("Tồn theo vòng đời đơn", "đơn mới → giữ tồn; huỷ hoặc hoàn trước khi giao → nhả tồn; giao hàng → trừ tồn thực và tính giá vốn. Tồn đẩy lên sàn = tồn thực − đang giữ − tồn an toàn."),
    ("Đẩy tồn", "mỗi khi tồn thay đổi, hệ thống gộp các thay đổi trong vài giây rồi tính lại số tồn cần hiển thị trên từng gian hàng và cập nhật lên sàn."),
    ("Đảm bảo", "chống bán vượt (khoá khi cập nhật tồn); combo trừ tất cả mã thành phần; mọi thay đổi tồn đều có dòng sổ ghi lại số dư sau thay đổi."),
])

h(doc, "4.4 Giao hàng & in vận đơn", 2)
flow(doc, "Tạo vận đơn → tải tem thật → in hàng loạt → quét đóng gói → bàn giao đơn vị vận chuyển → trừ tồn")
para(doc, "Hai cách giao hàng: (A) dùng đơn vị vận chuyển do sàn chỉ định — gọi sàn để sắp xếp vận chuyển rồi tải tem; "
          "(B) dùng đơn vị vận chuyển riêng của nhà bán (GHN, GHTK, J&T…) — tạo vận đơn và tải tem từ đơn vị đó.")
numbered(doc, [
    "Tạo vận đơn → nhận mã vận đơn và tem thật → lưu trữ an toàn (nếu sàn chưa kịp tạo tem, hệ thống tự thử lại sau).",
    "In hàng loạt: chọn nhiều đơn → hệ thống ghép tem thành một tệp để in một lần; phiếu soạn hàng và đóng gói được tự tạo.",
    "Quét đóng gói: nhân viên quét mã vận đơn để xác nhận đóng gói, rồi bàn giao cho đơn vị vận chuyển; khi bàn giao, đơn chuyển trạng thái đã giao và tồn được trừ.",
    "Phiếu in được lưu để in lại trong 90 ngày; sau đó tệp được tự xoá (chỉ giữ thông tin tối thiểu) để bảo vệ dữ liệu cá nhân.",
])
bullets(doc, [("Đảm bảo", "không tự vẽ lại tem của đơn vị vận chuyển; quét chống nhầm nhà bán và chống quét trùng.")])

h(doc, "4.5 Mua hàng → nhập kho → giá vốn", 2)
flow(doc, "Đề xuất nhập hàng → tạo đơn mua → nhận hàng → nhập kho (tăng tồn + lớp giá vốn) → giao hàng tính giá vốn FIFO")
numbered(doc, [
    "Hệ thống đề xuất số lượng cần nhập dựa trên tốc độ bán, thời gian đặt hàng, tồn hiện có và hàng đang về.",
    "Tạo đơn mua cho nhà cung cấp → xác nhận → nhận hàng → nhập kho.",
    "Nhập kho làm tăng tồn và tạo lớp giá vốn; khi giao hàng, hệ thống tính giá vốn theo nguyên tắc nhập trước-xuất trước (FIFO).",
])

h(doc, "4.6 Đối soát phí sàn → lợi nhuận thực", 2)
flow(doc, "Kéo bảng đối soát từ sàn → khớp từng dòng phí vào đơn → lợi nhuận theo phí thực tế")
numbered(doc, [
    "Hệ thống kéo bảng đối soát phí từ sàn (hoa hồng, phí thanh toán, phí vận chuyển, trợ giá, ưu đãi…).",
    "Khớp từng dòng phí vào đơn tương ứng.",
    "Lợi nhuận mỗi đơn = doanh thu − giá vốn − phí thực tế của sàn (không phải ước tính).",
])

h(doc, "4.7 Nhắn tin đa sàn & trợ lý AI", 2)
flow(doc, "Khách gửi tin → vào hộp thư hợp nhất → (tuỳ chọn) AI phân loại ý định → tin nhạy cảm chuyển nhân viên / tin an toàn AI tự trả lời")
numbered(doc, [
    "Khi khách gửi tin trên sàn hoặc Facebook, hệ thống tiếp nhận vào hộp thư hợp nhất (an toàn-khi-lặp-lại).",
    "Có thể trả lời tự động theo mẫu (chào lần đầu, theo từ khoá, theo lịch).",
    "Nếu nhà bán bật trợ lý AI và đủ điều kiện gói cước: hệ thống phân loại ý định tin nhắn trước. Tin nhạy cảm "
    "(khiếu nại, hoàn tiền, khẩn cấp) được chuyển cho nhân viên — AI không tự trả lời. Với tin an toàn, hệ thống "
    "che thông tin cá nhân, tham khảo tài liệu của shop, soạn câu trả lời và gửi đi; chi phí được ghi nhận.",
])
bullets(doc, [
    ("Đảm bảo", "mặc định an toàn (không chắc chắn thì chuyển cho người); nhà vận hành nền tảng quản lý nhà cung cấp AI, người bán chỉ chọn để dùng."),
])

h(doc, "4.8 Ghi sổ kế toán tự động", 2)
flow(doc, "Nghiệp vụ phát sinh (nhập kho/giao hàng/đối soát) → tự ghi bút toán nợ-có cân đối → báo cáo tài chính + xuất tệp")
bullets(doc, [
    ("Tự ghi sổ", "khi các nghiệp vụ phát sinh (nhập kho, giao hàng, đối soát phí…), hệ thống tự ghi bút toán kế toán kép tương ứng và đảm bảo cân đối nợ-có."),
    ("Kết quả", "hỗ trợ khoá kỳ kế toán, theo dõi công nợ phải thu/phải trả, lập báo cáo tài chính theo chuẩn mực kế toán Việt Nam (TT133), và xuất tệp tương thích phần mềm kế toán phổ biến (MISA)."),
])

h(doc, "4.9 Gói thuê bao & thanh toán", 2)
flow(doc, "Tạo nhà bán → dùng thử 14 ngày → vượt hạn mức thì gợi ý nâng cấp → thanh toán → kích hoạt gói")
numbered(doc, [
    "Nhà bán mới được tự dùng thử 14 ngày.",
    "Khi vượt hạn mức hoặc dùng tính năng ngoài gói, hệ thống chặn và gợi ý nâng cấp.",
    "Mua hoặc gia hạn qua chuyển khoản tự động (SePay) hoặc cổng thanh toán (VNPay) → kích hoạt gói.",
    "Hết hạn → có 7 ngày ân hạn rồi tự chuyển về gói dùng thử; hệ thống không khoá dữ liệu của nhà bán.",
])

# 5. USE CASES
h(doc, "5. Các trường hợp sử dụng chính", 1)
para(doc, "Mỗi trường hợp nêu: người dùng, điều kiện trước, luồng chính, kết quả.")
ucs = [
    ("UC-01 Kết nối gian hàng sàn", "Chủ shop / Quản trị",
     "Có tài khoản gian hàng trên sàn (TikTok / Shopee / Lazada).",
     "Vào mục Gian hàng → Kết nối → uỷ quyền trên sàn → hệ thống lưu kết nối và kéo lịch sử đơn 90 ngày.",
     "Gian hàng hoạt động; đơn cũ và mới bắt đầu đồng bộ về."),
    ("UC-02 Đơn mới tự đồng bộ và giữ tồn", "Hệ thống (tự động)",
     "Gian hàng đã kết nối.",
     "Nhận tín hiệu từ sàn hoặc đồng bộ định kỳ → lấy chi tiết đơn → lưu đơn → giữ tồn → khớp hồ sơ khách → ghi lịch sử trạng thái.",
     "Đơn xuất hiện trong vài phút, tồn đã được giữ, có hồ sơ khách kèm mức độ tín nhiệm."),
    ("UC-03 Tạo đơn thủ công", "Nhân viên xử lý đơn",
     "Có mã hàng gốc và tồn kho.",
     "Nhập người nhận và sản phẩm → tạo đơn tạo tay → giữ tồn ngay.",
     "Đơn đi vào luồng xử lý chung như đơn của sàn."),
    ("UC-04 Ghép mã hàng sàn với mã gốc", "Nhân viên xử lý đơn / Chủ shop",
     "Có sản phẩm trên sàn chưa ghép mã.",
     "Vào màn Liên kết mã hàng → ghép tự động theo mã trùng hoặc chọn tay (một-một hoặc combo) → xác nhận.",
     "Đơn chưa ghép tự được nối mã, giữ tồn và đẩy tồn lên sàn; hết cờ cần xử lý."),
    ("UC-05 Xử lý đơn → in tem → đóng gói → bàn giao", "Nhân viên xử lý đơn / kho",
     "Đơn đang xử lý, đủ tồn.",
     "Chuẩn bị hàng (chặn nếu thiếu tồn) → tạo vận đơn lấy tem → in tem hàng loạt → quét đóng gói → bàn giao đơn vị vận chuyển.",
     "Đơn chuyển trạng thái đã giao, trừ tồn; tem được lưu và in lại được trong 90 ngày."),
    ("UC-06 Quản lý kho", "Nhân viên kho",
     "Có kho và mã hàng.",
     "Tạo phiếu nhập / điều chuyển / kiểm kê (nháp) → xác nhận → ghi sổ tồn và giá vốn.",
     "Tồn cập nhật, có dấu vết kiểm toán, tự đẩy tồn lên sàn."),
    ("UC-07 Mua hàng theo đề xuất", "Chủ shop / Nhân viên",
     "Có lịch sử bán và nhà cung cấp.",
     "Xem đề xuất nhập hàng → tạo đơn mua theo nhà cung cấp → xác nhận → nhận hàng → nhập kho.",
     "Tồn tăng kèm lớp giá vốn; đơn mua chuyển trạng thái nhận một phần / đã nhận."),
    ("UC-08 Đối soát và lợi nhuận thực", "Kế toán / Chủ shop",
     "Đã bật đối soát cho sàn.",
     "Kéo bảng đối soát → khớp từng dòng phí vào đơn → tính lợi nhuận theo phí thực tế.",
     "Biết lãi/lỗ theo đơn, sản phẩm, gian hàng; cảnh báo dòng phí không khớp đơn."),
    ("UC-09 Kế toán & báo cáo tài chính", "Kế toán",
     "Gói có tính năng kế toán; đã thiết lập hệ thống tài khoản.",
     "Sổ tự ghi theo nghiệp vụ → xem công nợ → khoá kỳ → xuất báo cáo tài chính và tệp cho phần mềm kế toán.",
     "Sổ kép cân đối, công nợ rõ ràng, có tệp để nhập vào phần mềm kế toán bên ngoài."),
    ("UC-10 Chăm sóc khách hàng với AI", "Nhân viên chăm sóc khách hàng",
     "Gói có tính năng AI; quản trị hệ thống đã bật nhà cung cấp AI; nhà bán đã bật trợ lý và chế độ tự trả lời.",
     "Hộp thư hợp nhất → AI gợi ý hoặc tự trả lời tin an toàn; tin nhạy cảm chuyển cho nhân viên.",
     "Phản hồi nhanh; mọi câu trả lời được ghi nhật ký kèm chi phí."),
    ("UC-11 Mua / gia hạn gói", "Chủ shop",
     "Nhà bán đang dùng thử hoặc đã hết hạn.",
     "Chọn gói → thanh toán (SePay / VNPay) → kích hoạt gói.",
     "Mở hạn mức và tính năng theo gói; hết hạn thì về gói dùng thử, không mất dữ liệu."),
    ("UC-12 Vận hành nền tảng", "Quản trị hệ thống",
     "Đăng nhập khu vực quản trị.",
     "Cấu hình nhà cung cấp AI (kể cả nhà cung cấp tuỳ chỉnh), chỉnh gói và hạn mức, phát ưu đãi, gửi thông báo, tra cứu nhật ký toàn hệ thống.",
     "Thay đổi vận hành không cần triển khai lại phần mềm."),
    ("UC-13 Tuân thủ yêu cầu xoá dữ liệu người mua", "Hệ thống (tuân thủ)",
     "Nhận yêu cầu xoá dữ liệu từ sàn hoặc khi ngắt kết nối gian hàng.",
     "Ẩn danh hoá hồ sơ khách của gian hàng đó và xoá các tệp phiếu/ảnh liên quan; giữ lại số liệu thống kê không định danh.",
     "Đáp ứng yêu cầu bảo vệ dữ liệu cá nhân của sàn."),
]
for title, actor, pre, fl, post in ucs:
    h(doc, title, 2)
    table(doc, ["Khía cạnh", "Nội dung"], [
        ["Người dùng", actor], ["Điều kiện trước", pre], ["Luồng chính", fl], ["Kết quả", post],
    ], widths=[3.5, 13.5])

# 6. Nhóm dữ liệu (không nêu tên bảng kỹ thuật)
h(doc, "6. Các nhóm dữ liệu chính", 1)
para(doc, "Hệ thống tổ chức dữ liệu theo các nhóm nghiệp vụ sau (mô tả chức năng, không phải cấu trúc kỹ thuật):")
bullets(doc, [
    "Nhà bán & người dùng, phân quyền, nhật ký thao tác.",
    "Gian hàng kết nối & nhật ký đồng bộ.",
    "Đơn hàng, sản phẩm và mã hàng.",
    "Kho, tồn theo kho, sổ biến động tồn, giá vốn.",
    "Vận đơn và phiếu in.",
    "Hồ sơ khách hàng.",
    "Nhà cung cấp và đơn mua.",
    "Bảng đối soát phí và lợi nhuận.",
    "Sổ kế toán, công nợ, quỹ & ngân hàng, thuế.",
    "Hội thoại, mẫu trả lời, cấu hình và nhật ký trợ lý AI.",
    "Gói cước, hoá đơn và thanh toán.",
])

# 7. Tích hợp
h(doc, "7. Tích hợp & khả năng mở rộng", 1)
para(doc, "Mỗi loại nhà cung cấp là một mô-đun cắm-thêm độc lập; thêm nhà cung cấp mới không ảnh hưởng phần lõi.")
table(doc, ["Nhóm tích hợp", "Hiện có / Trạng thái"], [
    ["Sàn thương mại điện tử", "TikTok Shop, Lazada (đã tích hợp); đơn thủ công; Shopee đang chờ cấp quyền truy cập."],
    ["Đơn vị vận chuyển", "GHN và nhà bán tự quản; GHTK, J&T và các đơn vị khác bổ sung dần."],
    ["Cổng thanh toán", "SePay (chuyển khoản tự động) và VNPay; MoMo dự phòng."],
    ["Kênh nhắn tin", "Facebook, TikTok, Shopee, Lazada."],
    ["Mô hình AI", "Các mô hình AI lớn (Claude, GPT, Gemini, DeepSeek, OpenRouter…) và mô hình tuỳ chỉnh kết nối qua giao diện riêng của nhà bán/nhà vận hành."],
], widths=[4.5, 12.5])

# 8. Gói cước
h(doc, "8. Mô hình gói cước", 1)
table(doc, ["Gói", "Giá tháng / năm (VND)", "Số gian hàng", "Tính năng nổi bật"], [
    ["Dùng thử", "0 (14 ngày)", "2", "Tính năng cơ bản"],
    ["Starter", "99.000 / 990.000", "2", "Đồng bộ đơn, tồn kho, in tem, sổ khách hàng"],
    ["Pro", "199.000 / 1.990.000", "5", "Thêm mua hàng, giá vốn, lợi nhuận, đối soát, đề xuất nhập, kế toán nền tảng, hộp thư hợp nhất"],
    ["Business", "399.000 / 3.990.000", "10", "Thêm đăng bán đa sàn, tự động hoá, ưu tiên hỗ trợ, kế toán nâng cao, trợ lý AI nhắn tin"],
], widths=[2.6, 3.8, 2.4, 8.2])

# 9. Phi chức năng
h(doc, "9. Yêu cầu phi chức năng", 1)
bullets(doc, [
    ("Cách ly dữ liệu", "dữ liệu mỗi nhà bán được tách biệt; khu vực quản trị nền tảng tách riêng khỏi nhà bán."),
    ("Bảo mật & dữ liệu cá nhân", "mã hoá thông tin nhạy cảm; che thông tin cá nhân trước khi gửi cho AI; tuân thủ yêu cầu xoá dữ liệu người mua của từng sàn."),
    ("Độ tin cậy", "đồng bộ kép (tức thời và định kỳ); tự thử lại khi lỗi tạm thời; có cơ chế dừng gọi nhà cung cấp đang lỗi để tránh dồn lỗi; cho phép xem lại và chạy lại."),
    ("Hiệu năng & quy mô", "đáp ứng quy mô hàng trăm nghìn đơn mỗi tháng; xử lý nền và giới hạn tần suất gọi theo từng gian hàng."),
    ("Khả năng quan sát", "có nhật ký đồng bộ, giám sát lỗi và cảnh báo (gian hàng cần kết nối lại, tỉ lệ lỗi tăng…)."),
])

# 10. Phạm vi
h(doc, "10. Phạm vi & ngoài phạm vi", 1)
h(doc, "10.1 Trong phạm vi", 2)
para(doc, "Đầy đủ nghiệp vụ bán hàng đa sàn: đơn hàng, giao hàng và in ấn, kho, sản phẩm và đăng bán, mua hàng, "
          "tài chính và đối soát, báo cáo, hậu mãi, kế toán theo chuẩn Việt Nam, nhắn tin và trợ lý AI, cùng lớp "
          "dịch vụ nền tảng (nhiều nhà bán, gói cước, nhật ký).")
h(doc, "10.2 Ngoài phạm vi (giai đoạn này)", 2)
bullets(doc, [
    "Chỉ phục vụ thị trường Việt Nam — không đa quốc gia, không đa tiền tệ.",
    "Không tích hợp nguồn hàng quốc tế; chỉ nhà cung cấp trong nước.",
    "Không làm sàn thương mại điện tử riêng, không bán hàng tại quầy.",
    "Không làm tiếp thị / gửi tin quảng cáo hàng loạt.",
    "Hoá đơn điện tử: thực hiện ở giai đoạn sau.",
    "Không có ứng dụng di động riêng (có thể bổ sung công cụ quét đóng gói trên điện thoại sau).",
    "Không tự xây cổng thanh toán (đã tích hợp các cổng sẵn có).",
])

# 11. Mức độ hoàn thiện
h(doc, "11. Mức độ hoàn thiện", 1)
para(doc, "Các phân hệ cốt lõi đã hoạt động; một số phân hệ đang được mở rộng.")
table(doc, ["Phân hệ", "Trạng thái"], [
    ["Đồng bộ đơn đa sàn (TikTok, Lazada)", "Đã có"],
    ["Đơn thủ công, mã hàng, tồn kho", "Đã có"],
    ["Giao hàng & in vận đơn", "Đã có (đang bổ sung đơn vị vận chuyển và mẫu in)"],
    ["Quản lý kho & mua hàng & giá vốn", "Đã có"],
    ["Đối soát & báo cáo lợi nhuận", "Đã có"],
    ["Kế toán theo chuẩn Việt Nam", "Đã có"],
    ["Gói cước & thanh toán", "Đã có"],
    ["Nhắn tin & trợ lý AI", "Đang triển khai và mở rộng"],
    ["Shopee", "Chờ cấp quyền truy cập"],
], widths=[10, 7])

# 12. Phụ lục
h(doc, "12. Phụ lục: Thuật ngữ", 1)
table(doc, ["Thuật ngữ", "Ý nghĩa"], [
    ["Nhà bán (tenant)", "Một không gian làm việc độc lập của một người/đơn vị bán hàng."],
    ["Mã hàng gốc", "Đơn vị tồn kho gốc của nhà bán, là nguồn sự thật về tồn."],
    ["Sản phẩm trên sàn", "Sản phẩm hoặc biến thể được niêm yết trên một gian hàng của sàn."],
    ["Mô-đun tích hợp", "Thành phần kết nối tới một nhà cung cấp (sàn, vận chuyển, thanh toán, nhắn tin, AI); phần lõi chỉ làm việc với dữ liệu chuẩn hoá."],
    ["Trạng thái chuẩn", "Bộ trạng thái đơn dùng chung, ánh xạ từ trạng thái riêng của mỗi sàn."],
    ["Đối soát", "Bảng phí và tiền sàn trả về để tính lợi nhuận thực."],
    ["Giá vốn FIFO", "Giá vốn hàng bán theo nguyên tắc nhập trước-xuất trước."],
    ["Kiểm soát ý định", "Cơ chế phân loại tin nhắn để chặn AI tự trả lời các tin nhạy cảm."],
], widths=[4.5, 12.5])

set_update_fields(doc)
doc.save(OUT)

chk = Document(OUT)
print("OK ->", OUT)
print("paragraphs=%d tables=%d headings=%d size=%d bytes" % (
    len(chk.paragraphs), len(chk.tables),
    sum(1 for p in chk.paragraphs if p.style.name.startswith("Heading")),
    os.path.getsize(OUT)))
